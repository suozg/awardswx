#!/usr/bin/python3
# -*- coding: utf-8 -*-

import sys
import os
import re
import math

def check_inn(text):
    """
    Проверяет корректность 10-значного ИНН.
    Возвращает список ошибок и корректных ИНН.
    """
    results = {"errors": [], "valid": []}

    # Поиск 10-значных чисел
    matches = re.findall(r'\b\d{10}\b', str(text))

    for match in matches:
        digits = [int(char) for char in match]

        # Расчет контрольной суммы
        k1 = sum(x * y for x, y in zip(digits[:9], [-1, 5, 7, 9, 4, 6, 10, 5, 7]))
        k2 = k1 % 11
        checksum = 0 if k2 == 10 else k2

        # Сравнение контрольной суммы с последней цифрой
        if checksum == digits[9]:
            results["valid"].append(match)
        else:
            results["errors"].append(match)

    return results

def process_docx(file_path):
    import docx
    total_checked = 0
    total_errors = 0

    try:
        doc = docx.Document(file_path)
    except Exception as e:
        print(f"Помилка при відкритті файла: {e}")
        return

    for paragraph in doc.paragraphs:
        result = check_inn(paragraph.text)
        total_checked += len(result["valid"]) + len(result["errors"])
        total_errors += len(result["errors"])
        for error in result["errors"]:
            print(f"==> ERROR: {error}")
        for valid in result["valid"]:
            print(f"ok: {valid}")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                result = check_inn(cell.text)
                total_checked += len(result["valid"]) + len(result["errors"])
                total_errors += len(result["errors"])
                for error in result["errors"]:
                    print(f"==> ERROR: {error}")
                for valid in result["valid"]:
                    print(f"ok: {valid}")

    print(f"Всего ИНН: {total_checked}, з помилками: {total_errors}")

def process_xls(file_path):
    import xlrd
    try:
        workbook = xlrd.open_workbook(file_path)
    except Exception as e:
        print(f"Помилка при відкритті файла: {e}")
        return

    sheet = workbook.sheet_by_index(0)
    total_checked = 0
    total_errors = 0

    for row in range(sheet.nrows):
        for col in range(sheet.ncols):
            cell_value = sheet.cell_value(row, col)
            if not cell_value:
                continue
            result = check_inn(cell_value)
            total_checked += len(result["valid"]) + len(result["errors"])
            total_errors += len(result["errors"])
            for error in result["errors"]:
                print(f"==> ERROR: {error}")
            for valid in result["valid"]:
                print(f"ok: {valid}")

    print(f"Усього РНКПО: {total_checked}, Помилок: {total_errors}")

def process_xlsx(file_path):
    import openpyxl
    try:
        workbook = openpyxl.load_workbook(file_path)
    except Exception as e:
        print(f"Помилка при відкритті файла: {e}")
        return

    sheet = workbook.active
    total_checked = 0
    total_errors = 0

    for row in sheet.iter_rows():
        for cell in row:
            cell_value = cell.value
            if cell_value is None:
                continue
            result = check_inn(cell_value)
            total_checked += len(result["valid"]) + len(result["errors"])
            total_errors += len(result["errors"])
            for error in result["errors"]:
                print(f"==> ERROR: {error}")
            for valid in result["valid"]:
                print(f"ok: {valid}")

    print(f"Усього РНКПО: {total_checked}, Помилок: {total_errors}")

def main():
    if len(sys.argv) != 2:
        print(f"Використання: {__file__} <файл>")
        sys.exit(1)

    file_path = sys.argv[1]
    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".docx" or extension == ".doc":
        process_docx(file_path)
    elif extension == ".xls":
        process_xls(file_path)
    elif extension == ".xlsx":
        process_xlsx(file_path)
    else:
        print("Не той файл (тільки doc, docx, xls, xlsx).")
        sys.exit(1)

if __name__ == "__main__":
    main()
